"""
Enhanced Markdown Export Service

An enhanced version that properly handles Markdown formatting, images,
tables, and generates professional-looking documents.
"""

import uuid
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional, Callable
import logging
import markdown
from PIL import Image, ImageDraw, ImageFont
import io

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib import colors

logger = logging.getLogger(__name__)


class EnhancedMarkdownExportService:
    """Enhanced markdown export service with proper formatting and image support."""
    
    def __init__(self, output_dir: str = "exports"):
        """Initialize the enhanced markdown export service.
        
        Args:
            output_dir: Base output directory for exported files
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize markdown parser with extensions
        self.md = markdown.Markdown(extensions=[
            'tables', 'fenced_code', 'codehilite', 'toc', 'attr_list'
        ])
        
        # ReportLab styles
        self.styles = getSampleStyleSheet()
        self._setup_enhanced_reportlab_styles()
    
    def _setup_enhanced_reportlab_styles(self):
        """Setup enhanced ReportLab styles."""
        # Enhanced title style
        self.styles.add(ParagraphStyle(
            name='EnhancedTitle',
            parent=self.styles['Title'],
            fontSize=28,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=HexColor('#2c3e50'),
            fontName='Helvetica-Bold'
        ))
        
        # Enhanced heading styles
        for i in range(1, 7):
            self.styles.add(ParagraphStyle(
                name='EnhancedHeading{}'.format(i),
                parent=self.styles['Heading{}'.format(i)],
                fontSize=24 - (i * 2),
                spaceAfter=20 - (i * 2),
                textColor=HexColor('#34495e'),
                fontName='Helvetica-Bold'
            ))
        
        # Enhanced normal text style
        self.styles.add(ParagraphStyle(
            name='EnhancedNormal',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            textColor=HexColor('#2c3e50'),
            fontName='Helvetica',
            wordWrap='CJK'
        ))
        
        # Bold text style
        self.styles.add(ParagraphStyle(
            name='EnhancedBold',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            textColor=HexColor('#2c3e50'),
            fontName='Helvetica-Bold',
            wordWrap='CJK'
        ))
    
    def _parse_markdown_enhanced(self, markdown_content: str) -> List[Dict[str, Any]]:
        """Parse markdown content into structured elements with enhanced formatting."""
        elements = []
        
        # Split into lines
        lines = markdown_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                elements.append({'type': 'spacer', 'content': ''})
                i += 1
                continue
            
            # Headers
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                content = line.lstrip('#').strip()
                elements.append({
                    'type': 'heading{}'.format(level),
                    'content': content,
                    'level': level
                })
                i += 1
                continue
            
            # Images
            image_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if image_match:
                alt_text = image_match.group(1)
                image_path = image_match.group(2)
                elements.append({
                    'type': 'image',
                    'alt_text': alt_text,
                    'image_path': image_path,
                    'content': line
                })
                i += 1
                continue
            
            # Code blocks
            if line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if i < len(lines):
                    i += 1  # Skip closing ```
                
                elements.append({
                    'type': 'code_block',
                    'content': '\n'.join(code_lines),
                    'language': line[3:].strip() if len(line) > 3 else ''
                })
                continue
            
            # Lists
            if line.startswith(('- ', '* ', '+ ')) or re.match(r'^\d+\.', line):
                list_items = []
                while i < len(lines) and (lines[i].strip().startswith(('- ', '* ', '+ ')) or 
                                        re.match(r'^\d+\.', lines[i].strip())):
                    list_items.append(lines[i].strip())
                    i += 1
                
                elements.append({
                    'type': 'list',
                    'content': list_items
                })
                continue
            
            # Tables
            if '|' in line and line.count('|') >= 2:
                table_rows = []
                while i < len(lines) and '|' in lines[i] and lines[i].count('|') >= 2:
                    # Skip separator rows (|----|)
                    if not re.match(r'^\s*\|[\s\-:|]+\|\s*$', lines[i]):
                        table_rows.append(lines[i].strip())
                    i += 1
                
                if table_rows:
                    elements.append({
                        'type': 'table',
                        'content': table_rows
                    })
                continue
            
            # Regular paragraphs with enhanced formatting
            paragraph_lines = []
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(('#', '```', '- ', '* ', '+ ')) and '|' not in lines[i]:
                paragraph_lines.append(lines[i])
                i += 1
            
            if paragraph_lines:
                content = '\n'.join(paragraph_lines)
                # Check for bold text
                if '**' in content:
                    elements.append({
                        'type': 'paragraph_bold',
                        'content': content
                    })
                else:
                    elements.append({
                        'type': 'paragraph',
                        'content': content
                    })
                continue
            
            i += 1
        
        return elements
    
    def _get_actual_diagram_image(self, alt_text: str) -> Optional[bytes]:
        """Get the actual generated diagram image."""
        # Map alt text to diagram filenames
        diagram_mapping = {
            "DIA3 System Architecture": "dia3_system_architecture.png",
            "Intelligence Question Framework Process": "intelligence_framework_process.png",
            "Framework Categories Overview": "framework_categories_overview.png",
            "Monte Carlo Simulation Process": "monte_carlo_simulation_process.png",
            "Threat Evolution Forecasting": "threat_evolution_forecasting.png",
            "Art of War Integration Framework": "art_of_war_integration.png",
            "Risk Timeline Forecasting": "risk_timeline_forecasting.png",
            "Strategic Position Forecasting": "strategic_position_forecasting.png",
            "Intelligence Fusion Process": "intelligence_fusion_process.png",
            "Predictive Intelligence Forecasting": "predictive_intelligence_forecasting.png",
            "Capability Evolution Forecasting": "capability_evolution_forecasting.png",
            "Technology Adoption Forecasting": "technology_adoption_forecasting.png",
            "Alliance Dynamics Forecasting": "alliance_dynamics_forecasting.png",
            "Predictive Analysis Timeline": "predictive_analysis_timeline.png",
            "Decision Tree Analysis": "decision_tree_analysis.png",
            "Risk Assessment Matrix": "risk_assessment_matrix.png"
        }
        
        # Try to find the actual diagram
        diagram_filename = diagram_mapping.get(alt_text)
        if diagram_filename:
            # Look in multiple possible locations
            possible_paths = [
                Path("docs/white_papers/images") / diagram_filename,
                Path("temp/mermaid") / diagram_filename,
                Path("docs/white_papers/generated_pdfs/images") / diagram_filename
            ]
            
            for path in possible_paths:
                if path.exists():
                    try:
                        with open(path, 'rb') as f:
                            return f.read()
                    except Exception as e:
                        logger.warning(f"Could not read diagram {path}: {e}")
        
        return None
    
    def _generate_placeholder_image(self, alt_text: str, width: int = 400, height: int = 200) -> bytes:
        """Generate a placeholder image for missing diagrams."""
        # First try to get the actual diagram
        actual_image = self._get_actual_diagram_image(alt_text)
        if actual_image:
            return actual_image
        
        # Fall back to placeholder if actual diagram not found
        img = Image.new('RGB', (width, height), color='#f8f9fa')
        draw = ImageDraw.Draw(img)
        
        # Try to use a default font
        try:
            font = ImageFont.truetype("arial.ttf", 16)
        except:
            font = ImageFont.load_default()
        
        # Draw border
        draw.rectangle([0, 0, width-1, height-1], outline='#dee2e6', width=2)
        
        # Draw text
        text = "Diagram: {}".format(alt_text)
        bbox = draw.textbbox((0, 0), text, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        
        x = (width - text_width) // 2
        y = (height - text_height) // 2
        
        draw.text((x, y), text, fill='#6c757d', font=font)
        
        # Convert to bytes
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        return img_bytes.getvalue()
    
    def _process_bold_text(self, text: str) -> str:
        """Process bold text formatting for PDF (ReportLab)."""
        # Replace bold segments, but skip ones starting with Narrative
        def replace_bold(match: re.Match) -> str:
            inner = match.group(1)
            if inner.strip().startswith('Narrative'):
                return inner
            return '<b>{}</b>'.format(inner)

        # Process bold formatting
        text = re.sub(r'\*\*([^*]+)\*\*', replace_bold, text)
        
        # Process italic formatting
        def replace_italic(match: re.Match) -> str:
            inner = match.group(1)
            return '<i>{}</i>'.format(inner)
        
        text = re.sub(r'\*([^*]+)\*', replace_italic, text)
        
        return text
    
    async def export_markdown_to_pdf_enhanced(self,
                                           markdown_content: str,
                                           output_filename: Optional[str] = None,
                                           progress_callback: Optional[Callable] = None) -> Dict[str, Any]:
        """Export markdown content to PDF with enhanced formatting."""
        operation_id = str(uuid.uuid4())
        
        try:
            if progress_callback:
                progress_callback(10, "Initializing enhanced PDF export...")
            
            # Generate output filename
            if not output_filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = "export_enhanced_{}.pdf".format(timestamp)
            
            output_path = self.output_dir / output_filename
            
            if progress_callback:
                progress_callback(20, "Parsing markdown content...")
            
            # Parse markdown with enhanced formatting
            elements = self._parse_markdown_enhanced(markdown_content)
            
            if progress_callback:
                progress_callback(40, "Generating enhanced PDF document...")
            
            # Create PDF
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=A4,
                rightMargin=0.75*inch,
                leftMargin=0.75*inch,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch
            )
            
            story = []
            figure_index = 1
            
            figure_index = 1
            figure_index = 1
            figure_index = 1
            for i, element in enumerate(elements):
                if progress_callback:
                    progress = 40 + (i / len(elements)) * 50
                    progress_callback(progress, "Processing element {}/{}".format(i+1, len(elements)))
                
                if element['type'] == 'spacer':
                    story.append(Spacer(1, 12))
                
                elif element['type'].startswith('heading'):
                    level = element['level']
                    style_name = 'EnhancedHeading{}'.format(level) if level <= 6 else 'EnhancedHeading6'
                    story.append(Paragraph(element['content'], self.styles[style_name]))
                
                elif element['type'] == 'paragraph':
                    content = self._process_bold_text(element['content'])
                    story.append(Paragraph(content, self.styles['EnhancedNormal']))
                    story.append(Spacer(1, 12))
                
                elif element['type'] == 'paragraph_bold':
                    # Use normal paragraph style with inline bold tags to avoid bolding entire line
                    content = self._process_bold_text(element['content'])
                    story.append(Paragraph(content, self.styles['EnhancedNormal']))
                    story.append(Spacer(1, 12))
                
                elif element['type'] == 'image':
                    # Generate placeholder image
                    img_data = self._generate_placeholder_image(element['alt_text'])
                    img_stream = io.BytesIO(img_data)
                    
                    # Add image to PDF
                    img = RLImage(img_stream, width=4*inch, height=2*inch)
                    story.append(img)
                    story.append(Spacer(1, 12))
                    
                    # Add caption
                    if element['alt_text'] is not None:
                        raw_caption = element['alt_text'].strip()
                        # Remove leading 'Narrative:' if present (case-insensitive)
                        remaining = re.sub(r'^\s*Narrative:\s*', '', raw_caption, flags=re.IGNORECASE)
                        final_caption = (
                            "Figure {}: {}".format(figure_index, remaining)
                            if remaining else "Figure {}".format(figure_index)
                        )
                        figure_index += 1

                        caption = Paragraph("<i>{}</i>".format(final_caption),
                                          self.styles['EnhancedNormal'])
                        story.append(caption)
                        story.append(Spacer(1, 12))
                
                elif element['type'] == 'code_block':
                    # Enhanced code block formatting
                    code_text = "<code>{}</code>".format(element['content'])
                    story.append(Paragraph(code_text, self.styles['Code']))
                    story.append(Spacer(1, 12))
                
                elif element['type'] == 'list':
                    for item in element['content']:
                        # Process bold text in list items
                        processed_item = self._process_bold_text(item[2:])
                        story.append(Paragraph("• {}".format(processed_item), 
                                             self.styles['EnhancedNormal']))
                    story.append(Spacer(1, 12))
                
                elif element['type'] == 'table':
                    # Enhanced table formatting with proper text wrapping
                    table_rows_raw = []
                    for row in element['content']:
                        cells = [cell.strip() for cell in row.split('|')[1:-1]]
                        table_rows_raw.append(cells)

                    if table_rows_raw:
                        # Calculate available width (A4 page width minus margins)
                        available_width = A4[0] - (1.5 * inch)  # 0.75 inch margins on each side

                        num_cols = len(table_rows_raw[0]) if table_rows_raw else 1
                        col_width = available_width / max(1, num_cols)

                        # Convert each cell to Paragraph to ensure wrapping
                        table_data: List[List[Paragraph]] = []
                        for row_cells in table_rows_raw:
                            wrapped_cells = []
                            for cell_text in row_cells:
                                # Insert zero-width spaces to encourage wrapping in long tokens
                                safe_text = re.sub(r'([/_.-])', lambda m: m.group(1) + '\u200b', cell_text)
                                processed = self._process_bold_text(safe_text)
                                wrapped_cells.append(Paragraph(processed, self.styles['EnhancedNormal']))
                            table_data.append(wrapped_cells)

                        table = Table(table_data, colWidths=[col_width] * num_cols, repeatRows=1)
                        table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, 0), 10),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                            ('TOPPADDING', (0, 0), (-1, -1), 4),
                            ('LEFTPADDING', (0, 0), (-1, -1), 4),
                            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                        ]))

                        story.append(table)
                        story.append(Spacer(1, 12))
            
            if progress_callback:
                progress_callback(90, "Building enhanced PDF document...")
            
            # Build PDF
            doc.build(story)
            
            if progress_callback:
                progress_callback(100, "Enhanced PDF export completed")
            
            return {
                "success": True,
                "operation_id": operation_id,
                "output_path": str(output_path),
                "file_size": output_path.stat().st_size,
                "message": "Enhanced PDF export completed successfully"
            }
        
        except Exception as e:
            logger.error("Error in enhanced PDF export: {}".format(e))
            return {
                "success": False,
                "operation_id": operation_id,
                "error": str(e)
            }
    
    async def export_markdown_to_word_enhanced(self,
                                            markdown_content: str,
                                            output_filename: Optional[str] = None,
                                            progress_callback: Optional[Callable] = None) -> Dict[str, Any]:
        """Export markdown content to Word document with enhanced formatting."""
        operation_id = str(uuid.uuid4())
        
        try:
            if progress_callback:
                progress_callback(10, "Initializing enhanced Word export...")
            
            # Generate output filename
            if not output_filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = "export_enhanced_{}.docx".format(timestamp)
            
            output_path = self.output_dir / output_filename
            
            if progress_callback:
                progress_callback(20, "Parsing markdown content...")
            
            # Parse markdown with enhanced formatting
            elements = self._parse_markdown_enhanced(markdown_content)
            
            if progress_callback:
                progress_callback(40, "Creating enhanced Word document...")
            
            # Create Word document
            doc = Document()
            
            # Set up enhanced styles
            self._setup_enhanced_word_styles(doc)
            
            figure_index = 1
            for i, element in enumerate(elements):
                if progress_callback:
                    progress = 40 + (i / len(elements)) * 50
                    progress_callback(progress, "Processing element {}/{}".format(i+1, len(elements)))
                
                if element['type'] == 'spacer':
                    doc.add_paragraph()
                
                elif element['type'].startswith('heading'):
                    level = element['level']
                    # Use paragraph with heading style to allow inline formatting
                    style_name = 'Heading {}'.format(level if level in (1, 2, 3, 4) else 4)
                    para = doc.add_paragraph(style=style_name)
                    self._add_formatted_text(para, element['content'])
                
                elif element['type'] == 'paragraph':
                    para = doc.add_paragraph()
                    self._add_formatted_text(para, element['content'])
                
                elif element['type'] == 'paragraph_bold':
                    para = doc.add_paragraph()
                    self._add_formatted_text(para, element['content'])
                
                elif element['type'] == 'image':
                    # Generate placeholder image
                    img_data = self._generate_placeholder_image(element['alt_text'])
                    img_stream = io.BytesIO(img_data)
                    
                    # Add image to document
                    doc.add_picture(img_stream, width=Inches(4))
                    
                    # Add caption
                    if element['alt_text'] is not None:
                        raw_caption = element['alt_text'].strip()
                        remaining = re.sub(r'^\s*Narrative:\s*', '', raw_caption, flags=re.IGNORECASE)
                        final_caption = (
                            "Figure {}: {}".format(figure_index, remaining)
                            if remaining else "Figure {}".format(figure_index)
                        )
                        figure_index += 1

                        caption = doc.add_paragraph(final_caption)
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                elif element['type'] == 'code_block':
                    # Enhanced code block
                    code_para = doc.add_paragraph(element['content'])
                    code_para.style = 'No Spacing'
                
                elif element['type'] == 'list':
                    for item in element['content']:
                        list_para = doc.add_paragraph(style='List Bullet')
                        self._add_formatted_text(list_para, item[2:])
                
                elif element['type'] == 'table':
                    # Enhanced table formatting
                    table_data = []
                    for row in element['content']:
                        cells = [cell.strip() for cell in row.split('|')[1:-1]]
                        table_data.append(cells)
                    
                    if table_data:
                        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                        table.style = 'Table Grid'
                        
                        for i, row_data in enumerate(table_data):
                            row = table.rows[i]
                            for j, cell_data in enumerate(row_data):
                                cell = row.cells[j]
                                cell.text = ''
                                self._add_formatted_text(cell.paragraphs[0], cell_data)
            
            if progress_callback:
                progress_callback(90, "Saving enhanced Word document...")
            
            # Save document
            doc.save(str(output_path))
            
            if progress_callback:
                progress_callback(100, "Enhanced Word export completed")
            
            return {
                "success": True,
                "operation_id": operation_id,
                "output_path": str(output_path),
                "file_size": output_path.stat().st_size,
                "message": "Enhanced Word export completed successfully"
            }
        
        except Exception as e:
            logger.error("Error in enhanced Word export: {}".format(e))
            return {
                "success": False,
                "operation_id": operation_id,
                "error": str(e)
            }
    
    def _setup_enhanced_word_styles(self, doc: Document):
        """Setup enhanced Word document styles."""
        # Add custom styles for better formatting
        styles = doc.styles
        
        # Enhanced normal style
        if 'Enhanced Normal' not in styles:
            enhanced_normal = styles.add_style('Enhanced Normal', WD_STYLE_TYPE.PARAGRAPH)
            enhanced_normal.font.name = 'Calibri'
            enhanced_normal.font.size = Pt(11)
        
        # Enhanced heading styles
        for i in range(1, 4):
            style_name = 'Enhanced Heading {}'.format(i)
            if style_name not in styles:
                enhanced_heading = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                enhanced_heading.font.name = 'Calibri'
                enhanced_heading.font.size = Pt(16 - (i * 2))
                enhanced_heading.font.bold = True
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add formatted text to a Word paragraph."""
        # Process bold text, but exclude "Narrative"
        # First, temporarily replace "Narrative" to protect it
        text = text.replace("**Narrative**", "___NARRATIVE_PROTECTED___")
        text = text.replace("**Narrative", "___NARRATIVE_START_PROTECTED___")
        text = text.replace("Narrative**", "___NARRATIVE_END_PROTECTED___")
        
        # Process both bold and italic formatting
        # Split by both bold and italic patterns
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                bold_text = part[2:-2]
                if bold_text.strip().startswith('Narrative'):
                    # Don't bold Narrative
                    paragraph.add_run(bold_text)
                else:
                    run = paragraph.add_run(bold_text)
                    run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                # Italic text
                italic_text = part[1:-1]
                run = paragraph.add_run(italic_text)
                run.italic = True
            else:
                # Regular text
                if part.strip():
                    # Restore "Narrative" without bold formatting
                    part = part.replace("___NARRATIVE_PROTECTED___", "Narrative")
                    part = part.replace("___NARRATIVE_START_PROTECTED___", "Narrative")
                    part = part.replace("___NARRATIVE_END_PROTECTED___", "Narrative")
                    paragraph.add_run(part)
