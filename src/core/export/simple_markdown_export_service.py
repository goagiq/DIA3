"""
Simple Markdown Export Service

A simplified version that works without WeasyPrint dependencies.
Uses ReportLab for PDF and python-docx for Word documents.
"""

import asyncio
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional, Callable
import logging
import markdown
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

logger = logging.getLogger(__name__)


class SimpleMarkdownExportService:
    """Simplified markdown export service using ReportLab and python-docx."""
    
    def __init__(self, output_dir: str = "exports"):
        """Initialize the simple markdown export service.
        
        Args:
            output_dir: Base output directory for exported files
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize markdown parser
        self.md = markdown.Markdown(extensions=['tables', 'fenced_code', 'codehilite'])
        
        # ReportLab styles
        self.styles = getSampleStyleSheet()
        self._setup_reportlab_styles()
    
    def _setup_reportlab_styles(self):
        """Setup custom ReportLab styles."""
        # Custom title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Title'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=HexColor('#2c3e50')
        ))
        
        # Custom heading styles
        for i in range(1, 7):
            self.styles.add(ParagraphStyle(
                name=f'CustomHeading{i}',
                parent=self.styles[f'Heading{i}'],
                fontSize=20 - (i * 2),
                spaceAfter=15 - (i * 2),
                textColor=HexColor('#34495e')
            ))
    
    def _parse_markdown(self, markdown_content: str) -> List[Dict[str, Any]]:
        """Parse markdown content into structured elements."""
        elements = []
        
        # Split into lines
        lines = markdown_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                elements.append({'type': 'spacer', 'content': ''})
                i += 1
                continue
            
            # Headers
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                content = line.lstrip('#').strip()
                elements.append({
                    'type': f'heading{level}',
                    'content': content,
                    'level': level
                })
                i += 1
                continue
            
            # Code blocks
            if line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if i < len(lines):
                    i += 1  # Skip closing ```
                
                elements.append({
                    'type': 'code_block',
                    'content': '\n'.join(code_lines),
                    'language': line[3:].strip() if len(line) > 3 else ''
                })
                continue
            
            # Lists
            if line.startswith(('- ', '* ', '+ ')) or re.match(r'^\d+\.', line):
                list_items = []
                while i < len(lines) and (lines[i].strip().startswith(('- ', '* ', '+ ')) or 
                                        re.match(r'^\d+\.', lines[i].strip())):
                    list_items.append(lines[i].strip())
                    i += 1
                
                elements.append({
                    'type': 'list',
                    'content': list_items
                })
                continue
            
            # Tables
            if '|' in line and line.count('|') >= 2:
                table_rows = []
                while i < len(lines) and '|' in lines[i] and lines[i].count('|') >= 2:
                    # Skip separator rows (|----|)
                    if not re.match(r'^\s*\|[\s\-:|]+\|\s*$', lines[i]):
                        table_rows.append(lines[i].strip())
                    i += 1
                
                if table_rows:
                    elements.append({
                        'type': 'table',
                        'content': table_rows
                    })
                continue
            
            # Regular paragraphs
            paragraph_lines = []
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(('#', '```', '- ', '* ', '+ ')) and '|' not in lines[i]:
                paragraph_lines.append(lines[i])
                i += 1
            
            if paragraph_lines:
                elements.append({
                    'type': 'paragraph',
                    'content': '\n'.join(paragraph_lines)
                })
                continue
            
            i += 1
        
        return elements
    
    async def export_markdown_to_pdf(self,
                                   markdown_content: str,
                                   output_filename: Optional[str] = None,
                                   template_name: str = "whitepaper",
                                   custom_template: Optional[Dict[str, Any]] = None,
                                   progress_callback: Optional[Callable] = None) -> Dict[str, Any]:
        """Export markdown content to PDF using ReportLab."""
        operation_id = str(uuid.uuid4())
        
        try:
            if progress_callback:
                progress_callback(10, "Initializing PDF export...")
            
            # Generate output filename
            if not output_filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"export_{timestamp}.pdf"
            
            output_path = self.output_dir / output_filename
            
            if progress_callback:
                progress_callback(20, "Parsing markdown content...")
            
            # Parse markdown
            elements = self._parse_markdown(markdown_content)
            
            if progress_callback:
                progress_callback(40, "Generating PDF document...")
            
            # Create PDF
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=A4,
                rightMargin=1*inch,
                leftMargin=1*inch,
                topMargin=1*inch,
                bottomMargin=1*inch
            )
            
            story = []
            
            for i, element in enumerate(elements):
                if progress_callback:
                    progress = 40 + (i / len(elements)) * 50
                    progress_callback(progress, f"Processing element {i+1}/{len(elements)}")
                
                if element['type'] == 'spacer':
                    story.append(Spacer(1, 12))
                
                elif element['type'].startswith('heading'):
                    level = element['level']
                    style_name = f'CustomHeading{level}' if level <= 6 else 'CustomHeading6'
                    story.append(Paragraph(element['content'], self.styles[style_name]))
                
                elif element['type'] == 'paragraph':
                    story.append(Paragraph(element['content'], self.styles['Normal']))
                    story.append(Spacer(1, 12))
                
                elif element['type'] == 'code_block':
                    # Simple code block formatting
                    code_text = f"<code>{element['content']}</code>"
                    story.append(Paragraph(code_text, self.styles['Code']))
                    story.append(Spacer(1, 12))
                
                elif element['type'] == 'list':
                    for item in element['content']:
                        story.append(Paragraph(f"• {item[2:]}", self.styles['Normal']))
                    story.append(Spacer(1, 12))
                
                elif element['type'] == 'table':
                    # Simple table formatting
                    for row in element['content']:
                        story.append(Paragraph(row, self.styles['Normal']))
                    story.append(Spacer(1, 12))
            
            if progress_callback:
                progress_callback(90, "Building PDF document...")
            
            # Build PDF
            doc.build(story)
            
            if progress_callback:
                progress_callback(100, "PDF export completed")
            
            return {
                "success": True,
                "operation_id": operation_id,
                "output_path": str(output_path),
                "file_size": output_path.stat().st_size,
                "message": "PDF export completed successfully"
            }
        
        except Exception as e:
            logger.error(f"Error in PDF export: {e}")
            return {
                "success": False,
                "operation_id": operation_id,
                "error": str(e)
            }
    
    async def export_markdown_to_word(self,
                                    markdown_content: str,
                                    output_filename: Optional[str] = None,
                                    template_name: str = "whitepaper",
                                    custom_template: Optional[Dict[str, Any]] = None,
                                    progress_callback: Optional[Callable] = None) -> Dict[str, Any]:
        """Export markdown content to Word document using python-docx."""
        operation_id = str(uuid.uuid4())
        
        try:
            if progress_callback:
                progress_callback(10, "Initializing Word export...")
            
            # Generate output filename
            if not output_filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"export_{timestamp}.docx"
            
            output_path = self.output_dir / output_filename
            
            if progress_callback:
                progress_callback(20, "Parsing markdown content...")
            
            # Parse markdown
            elements = self._parse_markdown(markdown_content)
            
            if progress_callback:
                progress_callback(40, "Creating Word document...")
            
            # Create Word document
            doc = Document()
            
            for i, element in enumerate(elements):
                if progress_callback:
                    progress = 40 + (i / len(elements)) * 50
                    progress_callback(progress, f"Processing element {i+1}/{len(elements)}")
                
                if element['type'] == 'spacer':
                    doc.add_paragraph()
                
                elif element['type'].startswith('heading'):
                    level = element['level']
                    if level == 1:
                        heading = doc.add_heading(element['content'], level=1)
                    elif level == 2:
                        heading = doc.add_heading(element['content'], level=2)
                    elif level == 3:
                        heading = doc.add_heading(element['content'], level=3)
                    else:
                        heading = doc.add_heading(element['content'], level=4)
                
                elif element['type'] == 'paragraph':
                    doc.add_paragraph(element['content'])
                
                elif element['type'] == 'code_block':
                    # Add code block
                    code_para = doc.add_paragraph(element['content'])
                    code_para.style = 'No Spacing'
                
                elif element['type'] == 'list':
                    for item in element['content']:
                        doc.add_paragraph(item[2:], style='List Bullet')
                
                elif element['type'] == 'table':
                    # Simple table - just add as paragraphs for now
                    for row in element['content']:
                        doc.add_paragraph(row)
            
            if progress_callback:
                progress_callback(90, "Saving Word document...")
            
            # Save document
            doc.save(str(output_path))
            
            if progress_callback:
                progress_callback(100, "Word export completed")
            
            return {
                "success": True,
                "operation_id": operation_id,
                "output_path": str(output_path),
                "file_size": output_path.stat().st_size,
                "message": "Word export completed successfully"
            }
        
        except Exception as e:
            logger.error(f"Error in Word export: {e}")
            return {
                "success": False,
                "operation_id": operation_id,
                "error": str(e)
            }
    
    async def export_markdown_to_both(self,
                                    markdown_content: str,
                                    output_filename: Optional[str] = None,
                                    template_name: str = "whitepaper",
                                    custom_template: Optional[Dict[str, Any]] = None,
                                    progress_callback: Optional[Callable] = None) -> Dict[str, Any]:
        """Export markdown content to both PDF and Word documents."""
        operation_id = str(uuid.uuid4())
        
        try:
            # Generate base filename
            if not output_filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"export_{timestamp}"
            
            # Export to PDF
            pdf_result = await self.export_markdown_to_pdf(
                markdown_content,
                f"{output_filename}.pdf",
                template_name,
                custom_template,
                lambda progress, message: progress_callback(progress * 0.5, f"PDF: {message}") if progress_callback else None
            )
            
            # Export to Word
            word_result = await self.export_markdown_to_word(
                markdown_content,
                f"{output_filename}.docx",
                template_name,
                custom_template,
                lambda progress, message: progress_callback(50 + progress * 0.5, f"Word: {message}") if progress_callback else None
            )
            
            return {
                "success": pdf_result["success"] and word_result["success"],
                "operation_id": operation_id,
                "pdf_result": pdf_result,
                "word_result": word_result,
                "message": "Export to both formats completed"
            }
        
        except Exception as e:
            logger.error(f"Error in dual export: {e}")
            return {
                "success": False,
                "operation_id": operation_id,
                "error": str(e)
            }
