"""
Word Document Exporter

Generates Word documents from markdown content using python-docx with comprehensive styling support.
"""

import os
from pathlib import Path
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import logging

logger = logging.getLogger(__name__)


class WordExporter:
    """Generates Word documents from markdown content using python-docx."""
    
    def __init__(self):
        """Initialize Word exporter."""
        self.document = None
        self.styles = {}
    
    def export_to_word(self, 
                      markdown_elements: List[Any],
                      output_path: str,
                      template_config: Dict[str, Any],
                      images: Dict[str, str] = None,
                      progress_callback=None) -> bool:
        """Export markdown elements to Word document.
        
        Args:
            markdown_elements: List of parsed markdown elements
            output_path: Output Word file path
            template_config: Template configuration
            images: Dictionary mapping image references to file paths
            progress_callback: Optional progress callback function
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Create output directory if it doesn't exist
            output_dir = Path(output_path).parent
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Create new document
            self.document = Document()
            self._setup_document_styles(template_config)
            
            # Add document metadata
            self._add_document_metadata(template_config)
            
            # Process elements
            total_elements = len(markdown_elements)
            
            for i, element in enumerate(markdown_elements):
                if progress_callback:
                    progress = (i / total_elements) * 100
                    progress_callback(progress, f"Processing element {i+1}/{total_elements}")
                
                # Convert element to Word content
                self._element_to_word(element, template_config, images)
            
            # Add footer
            footer = self.document.sections[0].footer
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.text = "Copyright (c) 2025 Boeing Intelligence & Analytics All right reserved."
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_paragraph.style = self.document.styles['Normal']
            # Set footer font size and color
            for run in footer_paragraph.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(102, 102, 102)  # #666666
            
            # Save document
            if progress_callback:
                progress_callback(90, "Saving Word document...")
            
            self.document.save(output_path)
            
            if progress_callback:
                progress_callback(100, "Word document generation completed")
            
            logger.info(f"Successfully generated Word document: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error generating Word document: {e}")
            return False
    
    def _setup_document_styles(self, template_config: Dict[str, Any]):
        """Setup document styles based on template configuration.
        
        Args:
            template_config: Template configuration
        """
        word_styles = template_config.get('word_styles', {})
        
        # Get default styles
        styles = self.document.styles
        
        # Title style
        if 'title_style' in word_styles:
            title_style = styles[word_styles['title_style']]
        else:
            title_style = styles['Title']
        
        # Heading styles
        try:
            heading1_style = styles['Heading 1']
        except KeyError:
            heading1_style = styles['Title']
        
        try:
            heading2_style = styles['Heading 2']
        except KeyError:
            heading2_style = styles['Title']
        
        # Body style
        if 'body_style' in word_styles:
            body_style = styles[word_styles['body_style']]
        else:
            body_style = styles['Normal']
        
        # Store styles for later use
        self.styles = {
            'title': title_style,
            'heading1': heading1_style,
            'heading2': heading2_style,
            'body': body_style
        }
        
        # Apply custom font settings if specified
        if 'font_family' in word_styles:
            font_family = word_styles['font_family']
            for style in self.styles.values():
                if hasattr(style, 'font'):
                    style.font.name = font_family
        
        if 'font_size' in word_styles:
            font_size = word_styles['font_size']
            for style in self.styles.values():
                if hasattr(style, 'font'):
                    style.font.size = Pt(font_size)
    
    def _add_document_metadata(self, template_config: Dict[str, Any]):
        """Add document metadata.
        
        Args:
            template_config: Template configuration
        """
        metadata = template_config.get('metadata', {})
        
        # Set document properties
        core_props = self.document.core_properties
        if 'author' in metadata:
            core_props.author = metadata['author']
        if 'subject' in metadata:
            core_props.subject = metadata['subject']
        if 'keywords' in metadata:
            core_props.keywords = ', '.join(metadata['keywords'])
    
    def _element_to_word(self, 
                        element: Any, 
                        template_config: Dict[str, Any],
                        images: Dict[str, str] = None):
        """Convert a markdown element to Word content.
        
        Args:
            element: Markdown element
            template_config: Template configuration
            images: Dictionary mapping image references to file paths
        """
        try:
            element_type = element.element_type.value
            
            if element_type == 'header':
                self._header_to_word(element, template_config)
            elif element_type == 'paragraph':
                self._paragraph_to_word(element, template_config)
            elif element_type == 'list':
                self._list_to_word(element, template_config)
            elif element_type == 'table':
                self._table_to_word(element, template_config)
            elif element_type == 'code_block':
                self._code_block_to_word(element, template_config)
            elif element_type == 'mermaid':
                self._mermaid_to_word(element, template_config, images)
            elif element_type == 'image':
                self._image_to_word(element, template_config, images)
            elif element_type == 'blockquote':
                self._blockquote_to_word(element, template_config)
            elif element_type == 'horizontal_rule':
                self._horizontal_rule_to_word()
            elif element_type == 'text':
                self._text_to_word(element, template_config)
            elif element_type == 'link':
                self._link_to_word(element, template_config)
            else:
                logger.warning(f"Unsupported element type: {element_type}")
        
        except Exception as e:
            logger.error(f"Error converting element to Word: {e}")
    
    def _header_to_word(self, element: Any, template_config: Dict[str, Any]):
        """Convert header element to Word.
        
        Args:
            element: Header element
            template_config: Template configuration
        """
        level = element.level or 1
        text = element.content
        
        # Add header paragraph
        if level == 1:
            paragraph = self.document.add_paragraph(text, style=self.styles['title'])
        elif level == 2:
            paragraph = self.document.add_paragraph(text, style=self.styles['heading1'])
        else:
            paragraph = self.document.add_paragraph(text, style=self.styles['heading2'])
        
        # Apply custom styling if specified
        pdf_styles = template_config.get('pdf_styles', {})
        if level == 1 and 'title' in pdf_styles:
            self._apply_paragraph_styling(paragraph, pdf_styles['title'])
        elif level == 2 and 'heading1' in pdf_styles:
            self._apply_paragraph_styling(paragraph, pdf_styles['heading1'])
        elif 'heading2' in pdf_styles:
            self._apply_paragraph_styling(paragraph, pdf_styles['heading2'])
    
    def _paragraph_to_word(self, element: Any, template_config: Dict[str, Any]):
        """Convert paragraph element to Word.
        
        Args:
            element: Paragraph element
            template_config: Template configuration
        """
        text = element.content
        
        # Process inline elements
        processed_text = self._process_inline_elements(text, template_config)
        
        # Create paragraph with inline formatting
        paragraph = self._create_formatted_paragraph(processed_text, template_config)
        
        # Apply custom styling if specified
        pdf_styles = template_config.get('pdf_styles', {})
        if 'body' in pdf_styles:
            self._apply_paragraph_styling(paragraph, pdf_styles['body'])
    
    def _create_formatted_paragraph(self, text: str, template_config: Dict[str, Any], paragraph=None):
        """Create a paragraph with inline formatting.
        
        Args:
            text: Text with HTML tags
            template_config: Template configuration
            paragraph: Optional existing paragraph to use
        """
        import re
        
        # Create paragraph if not provided
        if paragraph is None:
            paragraph = self.document.add_paragraph(style=self.styles['body'])
        
        # Split text by HTML tags
        parts = re.split(r'(<[^>]+>.*?</[^>]+>)', text)
        
        for part in parts:
            if part.startswith('<b>') and part.endswith('</b>'):
                # Bold text
                bold_text = part[3:-4]  # Remove <b> and </b>
                run = paragraph.add_run(bold_text)
                run.bold = True
            elif part.startswith('<i>') and part.endswith('</i>'):
                # Italic text
                italic_text = part[3:-4]  # Remove <i> and </i>
                run = paragraph.add_run(italic_text)
                run.italic = True
            elif part.startswith('<code>') and part.endswith('</code>'):
                # Inline code
                code_text = part[6:-7]  # Remove <code> and </code>
                run = paragraph.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            elif part.startswith('<strike>') and part.endswith('</strike>'):
                # Strikethrough text
                strike_text = part[8:-9]  # Remove <strike> and </strike>
                run = paragraph.add_run(strike_text)
                run.font.strike = True
            else:
                # Regular text
                if part.strip():  # Only add non-empty parts
                    paragraph.add_run(part)
        
        return paragraph
    
    def _list_to_word(self, element: Any, template_config: Dict[str, Any]):
        """Convert list element to Word.
        
        Args:
            element: List element
            template_config: Template configuration
        """
        list_items = element.attributes.get('list_items', [])
        
        for item in list_items:
            marker = item['marker']
            item_text = item['content']
            
            # Remove numbering from the content (e.g., "1. " or "1.	" from the beginning)
            import re
            # Remove numbered list markers like "1. ", "2. ", etc. from the beginning
            cleaned_text = re.sub(r'^\d+\.\s*', '', item_text.strip())
            
            # Process inline markdown formatting in list items
            processed_text = self._process_inline_elements(cleaned_text, template_config)
            
            # Create paragraph with inline formatting (treat all lists as bullet points to remove numbering)
            paragraph = self._create_formatted_paragraph(processed_text, template_config)
            # Apply list bullet style for all lists (removes numbering)
            try:
                paragraph.style = 'List Bullet'
            except:
                # Fallback if style doesn't exist
                pass
    
    def _table_to_word(self, element: Any, template_config: Dict[str, Any]):
        """Convert table element to Word.
        
        Args:
            element: Table element
            template_config: Template configuration
        """
        table_data = element.attributes.get('table_data', {})
        headers = table_data.get('headers', [])
        data_rows = table_data.get('data', [])
        
        if not headers:
            return
        
        # Create table
        num_cols = len(headers)
        num_rows = len(data_rows) + 1  # +1 for header row
        
        table = self.document.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            # Process inline markdown formatting in headers
            processed_header = self._process_inline_elements(header, template_config)
            # Clear existing content and add formatted content
            header_cells[i].text = ""
            self._create_formatted_paragraph(processed_header, template_config, header_cells[i].paragraphs[0])
            # Make header bold
            for run in header_cells[i].paragraphs[0].runs:
                run.bold = True
        
        # Add data rows
        for i, row_data in enumerate(data_rows):
            row_cells = table.rows[i + 1].cells
            for j, cell_data in enumerate(row_data):
                if j < len(row_cells):
                    # Process inline markdown formatting in cell data
                    processed_cell_data = self._process_inline_elements(cell_data, template_config)
                    # Clear existing content and add formatted content
                    row_cells[j].text = ""
                    self._create_formatted_paragraph(processed_cell_data, template_config, row_cells[j].paragraphs[0])
    
    def _code_block_to_word(self, element: Any, template_config: Dict[str, Any]):
        """Convert code block element to Word.
        
        Args:
            element: Code block element
            template_config: Template configuration
        """
        code_text = element.content
        language = element.language or 'text'
        
        # Add code block paragraph
        paragraph = self.document.add_paragraph(code_text, style='No Spacing')
        
        # Apply code styling
        for run in paragraph.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        
        # Add spacing after code block
        self.document.add_paragraph()
    
    def _mermaid_to_word(self, 
                        element: Any, 
                        template_config: Dict[str, Any],
                        images: Dict[str, str] = None):
        """Convert Mermaid diagram element to Word.
        
        Args:
            element: Mermaid element
            template_config: Template configuration
            images: Dictionary mapping image references to file paths
        """
        # Try to find the corresponding image
        if images:
            # Look for the diagram image
            diagram_name = f"diagram_{len(self.document.paragraphs)}"  # Simple naming
            image_path = images.get(diagram_name)
            
            if image_path and os.path.exists(image_path):
                try:
                    # Add image to document
                    self.document.add_picture(image_path, width=Inches(6))
                    self.document.add_paragraph()  # Add spacing
                except Exception as e:
                    logger.warning(f"Could not add Mermaid diagram image: {e}")
                    # Fallback to code block
                    self._code_block_to_word(element, template_config)
            else:
                # Fallback to code block
                self._code_block_to_word(element, template_config)
        else:
            # Fallback to code block
            self._code_block_to_word(element, template_config)
    
    def _image_to_word(self, 
                      element: Any, 
                      template_config: Dict[str, Any],
                      images: Dict[str, str] = None):
        """Convert image element to Word.
        
        Args:
            element: Image element
            template_config: Template configuration
            images: Dictionary mapping image references to file paths
        """
        image_url = element.content
        alt_text = element.attributes.get('alt_text', '')
        
        # Try to find the image file
        image_path = None
        if images and image_url in images:
            image_path = images[image_url]
        elif os.path.exists(image_url):
            image_path = image_url
        
        if image_path and os.path.exists(image_path):
            try:
                # Add image to document
                self.document.add_picture(image_path, width=Inches(5))
                
                # Add alt text as caption
                if alt_text:
                    caption = self.document.add_paragraph(alt_text)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in caption.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(102, 102, 102)
                
                self.document.add_paragraph()  # Add spacing
            except Exception as e:
                logger.warning(f"Could not add image {image_path}: {e}")
                # Don't add any placeholder text - just skip the image
                pass
        else:
            # Don't add any placeholder text - just skip the image
            logger.warning(f"Image not found: {image_url}")
            pass
    
    def _blockquote_to_word(self, element: Any, template_config: Dict[str, Any]):
        """Convert blockquote element to Word.
        
        Args:
            element: Blockquote element
            template_config: Template configuration
        """
        text = element.content
        
        # Add blockquote paragraph
        paragraph = self.document.add_paragraph(text)
        
        # Apply blockquote styling
        paragraph.style = 'Quote'
        
        # Apply custom styling if specified
        pdf_styles = template_config.get('pdf_styles', {})
        if 'blockquote' in pdf_styles:
            self._apply_paragraph_styling(paragraph, pdf_styles['blockquote'])
    
    def _horizontal_rule_to_word(self):
        """Add horizontal rule to Word document."""
        # Add a paragraph with border
        paragraph = self.document.add_paragraph()
        paragraph.style = 'No Spacing'
        
        # Add border to paragraph
        try:
            p = paragraph._p
            pPr = p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom_border = OxmlElement('w:bottom')
            pBdr.append(bottom_border)
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '6')
            bottom_border.set(qn('w:space'), '1')
            bottom_border.set(qn('w:color'), 'auto')
            pPr.append(pBdr)
        except Exception as e:
            # Fallback: just add a simple horizontal line
            logger.warning(f"Could not add horizontal rule border: {e}")
            paragraph.add_run('_' * 50)  # Simple underline
    
    def _text_to_word(self, element: Any, template_config: Dict[str, Any]):
        """Convert text element to Word.
        
        Args:
            element: Text element
            template_config: Template configuration
        """
        text = element.content
        
        # Process inline elements
        processed_text = self._process_inline_elements(text, template_config)
        
        # Add text paragraph
        paragraph = self.document.add_paragraph(processed_text, style=self.styles['body'])
        
        # Apply custom styling if specified
        pdf_styles = template_config.get('pdf_styles', {})
        if 'body' in pdf_styles:
            self._apply_paragraph_styling(paragraph, pdf_styles['body'])
    
    def _apply_paragraph_styling(self, paragraph, style_config: Dict[str, Any]):
        """Apply custom styling to paragraph.
        
        Args:
            paragraph: Word paragraph object
            style_config: Style configuration dictionary
        """
        for run in paragraph.runs:
            if 'font_size' in style_config:
                run.font.size = Pt(style_config['font_size'])
            
            if 'color' in style_config:
                color = style_config['color']
                if color and color.startswith('#'):
                    # Convert hex to RGB
                    color = color.lstrip('#')
                    if len(color) == 6:  # Valid hex color
                        try:
                            r = int(color[0:2], 16)
                            g = int(color[2:4], 16)
                            b = int(color[4:6], 16)
                            run.font.color.rgb = RGBColor(r, g, b)
                        except ValueError:
                            # Skip invalid color values
                            pass
    
    def _process_inline_elements(self, text: str, template_config: Dict[str, Any]) -> str:
        """Process inline elements in text.
        
        Args:
            text: Text to process
            template_config: Template configuration
            
        Returns:
            Processed text with inline elements
        """
        import re
        
        # Process bold text (**text** or __text__)
        text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
        text = re.sub(r'__(.*?)__', r'<b>\1</b>', text)
        
        # Process italic text (*text* or _text_)
        text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
        text = re.sub(r'_(.*?)_', r'<i>\1</i>', text)
        
        # Process inline code (`code`)
        text = re.sub(r'`(.*?)`', r'<code>\1</code>', text)
        
        # Process strikethrough (~~text~~)
        text = re.sub(r'~~(.*?)~~', r'<strike>\1</strike>', text)
        
        return text
    
    def _link_to_word(self, element: Any, template_config: Dict[str, Any]):
        """Convert link element to Word.
        
        Args:
            element: Link element
            template_config: Template configuration
        """
        link_text = element.content
        link_url = element.attributes.get('url', '')
        
        # Add link paragraph
        paragraph = self.document.add_paragraph()
        
        if link_url:
            # Create hyperlink
            run = paragraph.add_run(f"{link_text} ({link_url})")
            run.font.color.rgb = RGBColor(52, 152, 219)  # Blue color
            run.font.underline = True
        else:
            # Just add text if no URL
            run = paragraph.add_run(link_text)
            run.font.color.rgb = RGBColor(52, 152, 219)  # Blue color
            run.font.underline = True
