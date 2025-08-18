#!/usr/bin/env python3
"""
Enhanced Word Document Generation Service for DIA3

This service provides Word document generation capabilities for white papers and reports
with mermaid diagram support and professional styling.
"""

import sys
import markdown
from pathlib import Path
import re
import base64
import requests
import hashlib
import asyncio
from typing import Dict, Any, Optional, List
from loguru import logger
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


class EnhancedWordGenerationService:
    """Enhanced service for generating Word documents from markdown content with mermaid diagrams."""
    
    def __init__(self):
        """Initialize the Word document generation service."""
        self.base_dir = Path(__file__).parent.parent.parent
        self.output_dir = self.base_dir / "docs" / "white_papers" / "generated_word"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info("✅ Enhanced Word Document Generation Service initialized")
    
    def generate_mermaid_image_sync(self, mermaid_code: str, diagram_type: str = "flowchart") -> Optional[bytes]:
        """Generate PNG image from mermaid code using mermaid.ink API (synchronous)."""
        try:
            # Encode the mermaid code for URL
            encoded_code = base64.b64encode(mermaid_code.encode()).decode()
            
            # Use mermaid.ink API to generate PNG
            url = f"https://mermaid.ink/img/{encoded_code}?type=png&theme=default"
            
            logger.info(f"  Generating {diagram_type} diagram...")
            
            # Download the image
            response = requests.get(url, timeout=30)
            if response.status_code == 200:
                return response.content
            else:
                logger.warning(f"  Warning: Failed to generate {diagram_type} diagram "
                              f"(HTTP {response.status_code})")
                return None
                
        except Exception as e:
            logger.warning(f"  Warning: Error generating {diagram_type} diagram: {e}")
            return None
    
    def convert_mermaid_to_images(self, markdown_content: str, output_dir: Path) -> tuple[str, List[Dict[str, Any]]]:
        """Convert mermaid diagrams to embedded images and return processed content with image info."""
        mermaid_pattern = r'```mermaid\s*\n(.*?)\n```'
        image_info_list = []
        
        def replace_mermaid(match):
            diagram_content = match.group(1)
            
            # Determine diagram type
            if 'flowchart' in diagram_content:
                diagram_type = "System Architecture"
                description = "This diagram shows the system architecture and data flow for the DIA3 Strategic Intelligence Framework."
            elif 'graph' in diagram_content:
                diagram_type = "Component Relationship"
                description = "This diagram illustrates the relationships between different system components."
            else:
                diagram_type = "Process Flow"
                description = "This diagram demonstrates the analysis process flow and decision points."
            
            # Generate image synchronously
            image_data = self.generate_mermaid_image_sync(diagram_content, diagram_type.lower())
            
            if image_data:
                # Save image to file
                content_hash = hashlib.md5(diagram_content.encode()).hexdigest()[:8]
                image_filename = f"diagram_{diagram_type.lower().replace(' ', '_')}_{content_hash}.png"
                image_path = output_dir / image_filename
                
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                
                # Store image information for Word document
                image_info = {
                    'path': str(image_path),
                    'filename': image_filename,
                    'type': diagram_type,
                    'description': description,
                    'data': image_data
                }
                image_info_list.append(image_info)
                
                # Return placeholder for markdown processing
                return f'[DIAGRAM_PLACEHOLDER_{len(image_info_list)-1}]'
            else:
                # Fallback to text description if image generation fails
                return f'\n**{diagram_type} Diagram**\n\n*Note: Diagram could not be generated automatically.*\n\n{description}\n'
        
        processed_content = re.sub(mermaid_pattern, replace_mermaid, markdown_content, flags=re.DOTALL)
        return processed_content, image_info_list
    
    def markdown_to_html(self, markdown_content: str) -> str:
        """Convert markdown to HTML with extensions."""
        extensions = [
            'markdown.extensions.tables',
            'markdown.extensions.fenced_code',
            'markdown.extensions.codehilite',
            'markdown.extensions.toc',
            'markdown.extensions.attr_list',
            'markdown.extensions.def_list',
            'markdown.extensions.footnotes',
        ]
        
        html = markdown.markdown(markdown_content, extensions=extensions)
        return html
    
    def create_word_document(self, title: str, content: str, image_info_list: List[Dict[str, Any]]) -> Document:
        """Create a Word document with content and embedded images."""
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = title
        doc.core_properties.author = "DIA3 Strategic Intelligence Framework"
        doc.core_properties.subject = "Strategic Intelligence Analysis"
        
        # Add title
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing after title
        doc.add_paragraph()
        
        # Process content and add to document
        lines = content.split('\n')
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            
            # Skip empty lines
            if not line:
                current_paragraph = None
                continue
            
            # Check for diagram placeholders
            if line.startswith('[DIAGRAM_PLACEHOLDER_'):
                placeholder_index = int(line.split('_')[2].rstrip(']'))
                if placeholder_index < len(image_info_list):
                    image_info = image_info_list[placeholder_index]
                    
                    # Add diagram title
                    diagram_title = doc.add_paragraph()
                    diagram_title_run = diagram_title.add_run(image_info['type'] + " Diagram")
                    diagram_title_run.font.size = Pt(14)
                    diagram_title_run.font.bold = True
                    diagram_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add image
                    try:
                        doc.add_picture(image_info['path'], width=Inches(6))
                        # Center the image
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        logger.warning(f"Could not add image {image_info['filename']}: {e}")
                        # Add fallback text
                        fallback_para = doc.add_paragraph()
                        fallback_para.add_run(f"[Image: {image_info['type']} Diagram]")
                        fallback_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add description
                    desc_para = doc.add_paragraph()
                    desc_para.add_run(image_info['description'])
                    desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    desc_para.style.font.italic = True
                    
                    # Add spacing
                    doc.add_paragraph()
                    current_paragraph = None
                    continue
            
            # Handle headers
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                header_text = line.lstrip('#').strip()
                
                if level == 1:
                    header_para = doc.add_paragraph()
                    header_run = header_para.add_run(header_text)
                    header_run.font.size = Pt(18)
                    header_run.font.bold = True
                elif level == 2:
                    header_para = doc.add_paragraph()
                    header_run = header_para.add_run(header_text)
                    header_run.font.size = Pt(16)
                    header_run.font.bold = True
                elif level == 3:
                    header_para = doc.add_paragraph()
                    header_run = header_para.add_run(header_text)
                    header_run.font.size = Pt(14)
                    header_run.font.bold = True
                else:
                    header_para = doc.add_paragraph()
                    header_run = header_para.add_run(header_text)
                    header_run.font.size = Pt(12)
                    header_run.font.bold = True
                
                current_paragraph = None
                continue
            
            # Handle lists
            if line.startswith('- ') or line.startswith('* '):
                list_text = line[2:].strip()
                list_para = doc.add_paragraph(list_text, style='List Bullet')
                current_paragraph = None
                continue
            
            if line.startswith('1. '):
                list_text = line[3:].strip()
                list_para = doc.add_paragraph(list_text, style='List Number')
                current_paragraph = None
                continue
            
            # Handle code blocks
            if line.startswith('```'):
                # Skip code block markers
                continue
            
            # Handle regular text
            if current_paragraph is None:
                current_paragraph = doc.add_paragraph()
            
            # Check for emphasis
            if '**' in line or '__' in line:
                # Handle bold text
                parts = re.split(r'(\*\*.*?\*\*|__.*?__)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        bold_text = part[2:-2]
                        current_paragraph.add_run(bold_text).bold = True
                    elif part.startswith('__') and part.endswith('__'):
                        bold_text = part[2:-2]
                        current_paragraph.add_run(bold_text).bold = True
                    elif part:
                        current_paragraph.add_run(part)
            else:
                current_paragraph.add_run(line + ' ')
        
        return doc
    
    async def generate_word_from_markdown(self, markdown_file: str, output_name: str = None, title: str = None) -> Dict[str, Any]:
        """Generate Word document from markdown file."""
        try:
            markdown_path = Path(markdown_file)
            if not markdown_path.exists():
                return {
                    "success": False,
                    "error": f"Markdown file not found: {markdown_file}"
                }
            
            # Determine output name
            if output_name is None:
                output_name = markdown_path.stem
            
            # Determine title
            if title is None:
                title = f"DIA3: {output_name.replace('_', ' ').title()}"
            
            # Generate output path
            output_file = self.output_dir / f"{output_name}.docx"
            
            # Read markdown content
            try:
                with open(markdown_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            except FileNotFoundError:
                logger.error(f"✗ File not found: {markdown_path}")
                return {
                    "success": False,
                    "error": f"File not found: {markdown_path}"
                }
            except Exception as e:
                logger.error(f"✗ Error reading file {markdown_path}: {e}")
                return {
                    "success": False,
                    "error": f"Error reading file: {e}"
                }
            
            # Create images directory
            images_dir = self.output_dir / "images"
            images_dir.mkdir(exist_ok=True)
            
            # Convert mermaid diagrams to images
            logger.info("Converting mermaid diagrams to images...")
            processed_content, image_info_list = self.convert_mermaid_to_images(content, images_dir)
            
            # Create Word document
            logger.info("Creating Word document...")
            doc = self.create_word_document(title, processed_content, image_info_list)
            
            # Save Word document
            try:
                doc.save(str(output_file))
                logger.success(f"✓ Successfully generated: {output_file}")
                
                return {
                    "success": True,
                    "word_file": str(output_file),
                    "title": title,
                    "diagrams_count": len(image_info_list),
                    "message": "Word document generated successfully with embedded diagrams."
                }
                
            except Exception as e:
                logger.error(f"✗ Error saving Word document: {e}")
                return {
                    "success": False,
                    "error": f"Error saving Word document: {e}"
                }
                
        except Exception as e:
            logger.error(f"Error generating Word document: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def generate_white_paper_word_documents(self) -> Dict[str, Any]:
        """Generate Word documents for all white papers."""
        try:
            docs_dir = self.base_dir / "docs" / "white_papers"
            
            # Define files to convert
            files_to_convert = [
                {
                    'input': docs_dir / "DIA3_Strategic_Intelligence_White_Paper.md",
                    'output_name': "DIA3_Strategic_Intelligence_White_Paper",
                    'title': "DIA3: Decision Intelligence Agentic, Autonomous, & Adaptive - Strategic Intelligence Framework White Paper"
                },
                {
                    'input': docs_dir / "DIA3_White_Paper_Summary.md",
                    'output_name': "DIA3_White_Paper_Executive_Summary",
                    'title': "DIA3 Strategic Intelligence Framework - Executive Summary"
                }
            ]
            
            results = []
            success_count = 0
            
            for file_info in files_to_convert:
                if file_info['input'].exists():
                    result = await self.generate_word_from_markdown(
                        str(file_info['input']),
                        file_info['output_name'],
                        file_info['title']
                    )
                    results.append(result)
                    if result['success']:
                        success_count += 1
                else:
                    results.append({
                        "success": False,
                        "error": f"File not found: {file_info['input']}"
                    })
            
            return {
                "success": success_count > 0,
                "total_files": len(files_to_convert),
                "successful_files": success_count,
                "results": results,
                "output_directory": str(self.output_dir)
            }
            
        except Exception as e:
            logger.error(f"Error generating white paper Word documents: {e}")
            return {
                "success": False,
                "error": str(e)
            }


# Global service instance
enhanced_word_service = EnhancedWordGenerationService()
